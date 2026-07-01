from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission"
OUT.mkdir(parents=True, exist_ok=True)
PDF = OUT / "MoonWheelKit_项目申报书_提交版.pdf"
DOCX = OUT / "MoonWheelKit_项目申报书_提交版.docx"
MD = OUT / "MoonWheelKit_项目申报书_提交版.md"
CHECKLIST = OUT / "MoonWheelKit_报名检查清单.md"

TITLE = "MoonWheelKit：面向 MoonBit 的分层时间轮与确定性虚拟时钟调度基础库"
REPO = "https://github.com/brother-666/MoonWheelKit"
ACCENT_HEX = "245B78"
ACCENT = colors.HexColor("#245B78")
INK = colors.HexColor("#202A33")
MUTED = colors.HexColor("#66737D")
LIGHT = colors.HexColor("#EDF4F7")

BASIC = [
    ("项目名称", TITLE),
    ("参赛者", "孙加豪"),
    ("联系方式", "【提交前填写：手机、邮箱】"),
    ("GitHub 仓库", REPO),
    ("GitLink 仓库", "https://www.gitlink.org.cn/brother666/MoonWheelKit"),
    ("项目方向", "MoonBit 基础生态库 / 定时调度 / 确定性测试基础设施"),
    ("是否为移植项目", "否"),
    ("当前基础", "15+ 次有效提交；21 项测试；四后端验证；GitHub CI 已通过"),
]

PAGE1 = [
    (
        "二、项目简介",
        [
            "MoonWheelKit 面向需要管理大量截止时间的 MoonBit 程序，提供可嵌入的分层时间轮和虚拟时钟调度内核。项目不读取系统时钟，不创建线程，也不绑定浏览器、Node.js 或操作系统接口；调用方通过推进逻辑 tick 驱动调度，因此同一操作序列在 Native、JavaScript、WebAssembly 和 Wasm-GC 后端具有一致的触发结果。",
            "项目适用于网关超时、缓存失效、重试队列、游戏逻辑、离散事件模拟、测试替身和异步运行时适配。它解决的不是单个任务如何 sleep，而是大量未来截止时间如何分层组织、低成本取消、迟到后如何追赶以及调度状态如何恢复。",
        ],
    ),
    (
        "三、拟解决的问题与生态价值",
        [
            "MoonBit 已有异步运行时和 sleep、timeout、retry 等能力，但这些接口面向真实事件循环，难以直接承担可重复的虚拟时间测试、批量截止时间管理和调度状态恢复。业务若自行维护排序数组或零散 timer，取消通常需要扫描，迟到与周期漂移语义也容易因后端而异。",
            "MoonWheelKit 把时间轮作为独立基础设施交付：上层可接入真实时钟、游戏帧或模拟时钟，核心始终保持无 IO、后端中立。库同时公开统计、校验、快照和 JSON 证据，便于测试框架、调试器和教学可视化复用。",
        ],
    ),
    (
        "四、核心功能范围",
        [
            "可配置层数、每层槽数、tick 单位和最大追赶次数的分层时间轮。",
            "单次、固定延迟、固定频率三类任务；相同截止时间保持稳定顺序。",
            "基于 timer id 与 generation 的 O(1) 表查找取消、重排和陈旧引用失效。",
            "逐 tick 精确推进，以及不扫描跳过区间的迟到推进和 lateness 报告。",
            "固定频率任务有界追赶，长时间暂停后不会产生无上限任务风暴。",
            "逻辑快照、确定性恢复、运行统计、JSON 输出、不变量校验和桶清理。",
        ],
    ),
    (
        "五、社区查重与边界",
        [
            "2026 年 6 月 30 日以 timing wheel、hierarchical timer、virtual time scheduler、cron、delayed task 等关键词检索 Mooncakes。moonbitlang/async 提供真实事件循环、sleep、timeout 和 retry；peter-jerry-ye/async 提供 promise、event loop、channel 与 IO；f4ah6o/mhx 的 delay、throttle、debounce 面向前端交互；dowdiness/moondsp 的调度面向音频领域。",
            "本项目不替代异步运行时、cron 解析器或持久化数据库。检索未发现以“分层槽位 + generation 失效 + 显式虚拟时间 + 有界迟到追赶 + 快照回放”为组合定位的通用 MoonBit 包。相关链接和差异逐项记录在 docs/RELATED_WORK.md。",
        ],
    ),
]

PAGE2 = [
    (
        "六、技术设计",
        [
            "设每层槽数为 S，第 L 层单槽跨度为 S^L tick。新任务进入能够覆盖其 deadline 差值的最低层；时间推进到层级边界时，高层槽中的有效引用逐级下放，临近任务最终进入第 0 层触发。",
            "任务真值只保存在中心 Map 中，槽位仅保存 timer id 与 generation。取消或重排只修改任务状态和 generation，无需在线性桶中查找旧节点；扫描时 generation 不匹配的引用会自然失效，compact 可重建所有有效桶。",
            "advance_to 用于精确逐 tick 观察；advance_late_to 直接跳到目标时间。固定延迟按“实际观察时间 + period”续排，固定频率按“原计划时间 + period”续排，并受 max_catch_up 限制。快照记录几何配置、时钟、计数器与任务状态，恢复后继续产生相同事件序列。",
        ],
    ),
    (
        "七、创新点",
        [
            "将高吞吐时间轮与可测试虚拟时间统一为一个后端中立内核，而不是事件循环内部不可复用的实现。",
            "generation 失效机制把取消、重排从槽位删除中解耦，并通过显式压缩和不变量校验给出可审计状态。",
            "同时定义精确推进和迟到推进，明确区分 FixedDelay 与 FixedRate 漂移，并用有界追赶处理恢复任务风暴。",
            "快照、恢复、稳定排序和四后端差异测试共同保证可重复性，可用于模拟器、测试框架和运行时适配。",
        ],
    ),
    (
        "八、当前成果与验证",
        [
            "仓库按功能形成 15 次有效提交，覆盖工程骨架、模型、时间轮、基础测试、级联与周期测试、诊断、快照、迟到语义、CLI、查重文档、万级工作负载、README、CI 和公开协作跟踪。",
            "21 项确定性测试覆盖精确截止时间、零延迟、稳定顺序、取消、重排、跨层级联、周期任务、迟到、有界追赶、压缩、统计、校验和快照恢复。默认、JS、Wasm、Wasm-GC 四组测试均通过，GitHub Actions 首次运行成功（run 28453414903）。",
            "bench/main 确定性调度 10,000 个任务并推进 4,096 tick，结果为 scheduled=10000、fired=10000、pending=0、validation_issues=0。仓库不虚构跨机器性能比例，后续按固定环境发布带日期的测量记录。",
        ],
    ),
    (
        "九、已知限制",
        [
            "当前时间值使用 Int tick，payload 为 String；快照已可导出 JSON，但尚未提供版本化 JSON 解析器。核心不主动读取真实时钟，也不提供线程安全承诺；这些能力应由适配层按目标后端实现。",
        ],
    ),
    (
        "十、后续计划",
        [
            "0.2：时钟适配 trait、批量调度与取消、版本化快照解析、模型对照随机测试。",
            "0.3：排序数组与二叉堆基线，四后端定期性能报告，聚集和长尾截止时间工作负载。",
            "0.4：合并窗口、任务组取消、逐任务追赶策略和可注入的确定性抖动策略。",
        ],
    ),
    (
        "十一、交付物与验收标准",
        [
            "源码交付包括公开 API、21 项测试、CLI 场景、万级确定性工作负载、四后端 CI、架构与语义文档、生态差异说明、变更日志和公开议题清单。",
            "阶段验收以可运行证据为准：moon fmt --check 与 moon check 无告警；四后端测试结果一致；快照恢复前后事件序列一致；万级工作负载无漏触发且校验结果为零。",
            "开发过程在公开仓库持续进行，以功能提交、Issue、Pull Request、CHANGELOG 和带日期的性能记录保持可追踪，不以一次性上传代替开源协作。",
        ],
    ),
]


def register_cn_font():
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("CN", str(path)))
            return "CN"
    return "Helvetica"


FONT = register_cn_font()


def escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont(FONT, 7.6)
    canvas.setFillColor(MUTED)
    canvas.drawString(16 * mm, 8 * mm, "MoonWheelKit 项目申报书")
    canvas.drawRightString(A4[0] - 16 * mm, 8 * mm, f"第 {doc.page} 页 / 共 2 页")
    canvas.restoreState()


def build_pdf():
    doc = BaseDocTemplate(
        str(PDF),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=12 * mm,
        bottomMargin=14 * mm,
    )
    doc.addPageTemplates(
        [
            PageTemplate(
                id="proposal",
                frames=[
                    Frame(
                        doc.leftMargin,
                        doc.bottomMargin,
                        doc.width,
                        doc.height,
                        id="content",
                    )
                ],
                onPage=pdf_footer,
            )
        ]
    )
    title = ParagraphStyle(
        "title",
        fontName=FONT,
        fontSize=14.4,
        leading=18,
        textColor=ACCENT,
        alignment=1,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "subtitle",
        fontName=FONT,
        fontSize=8.4,
        leading=10.5,
        textColor=MUTED,
        alignment=1,
        spaceAfter=5,
    )
    heading = ParagraphStyle(
        "heading",
        fontName=FONT,
        fontSize=10.5,
        leading=13,
        textColor=ACCENT,
        spaceBefore=3,
        spaceAfter=2,
    )
    body = ParagraphStyle(
        "body",
        fontName=FONT,
        fontSize=8.75,
        leading=12,
        textColor=INK,
        firstLineIndent=13,
        spaceAfter=1.8,
    )
    bullet = ParagraphStyle(
        "bullet",
        fontName=FONT,
        fontSize=8.65,
        leading=11.7,
        textColor=INK,
        leftIndent=11,
        firstLineIndent=-8,
        spaceAfter=1.4,
    )
    small = ParagraphStyle(
        "small",
        fontName=FONT,
        fontSize=7.9,
        leading=10.2,
        textColor=INK,
    )
    story = [
        Paragraph("MoonWheelKit 项目申报书", title),
        Paragraph("分层时间轮与确定性虚拟时钟调度基础库", subtitle),
        Paragraph("一、基本信息", heading),
    ]
    table = Table(
        [[Paragraph(escape(k), small), Paragraph(escape(v), small)] for k, v in BASIC],
        colWidths=[31 * mm, 131 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), LIGHT),
                ("TEXTCOLOR", (0, 0), (0, -1), ACCENT),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D9E1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 2.3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.3),
            ]
        )
    )
    story.extend([table, Spacer(1, 2)])
    for page_index, sections in enumerate([PAGE1, PAGE2]):
        if page_index == 1:
            story.append(PageBreak())
        for section_title, paragraphs in sections:
            story.append(Paragraph(escape(section_title), heading))
            use_bullets = section_title in {
                "四、核心功能范围",
                "七、创新点",
                "十、后续计划",
            }
            for index, text in enumerate(paragraphs):
                prefix = f"{index + 1}. " if use_bullets else ""
                story.append(Paragraph(escape(prefix + text), bullet if use_bullets else body))
    doc.build(story)
    pages = len(PdfReader(str(PDF)).pages)
    if pages != 2:
        raise RuntimeError(f"PDF page count must be 2, got {pages}")


def set_run_font(run, size=9.2, bold=False, color="202A33"):
    run.font.name = "Microsoft YaHei"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        fonts.set(qn(key), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_docx_sections(doc, sections):
    for section_title, paragraphs in sections:
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        set_run_font(heading.add_run(section_title), 10.8, True, ACCENT_HEX)
        use_bullets = section_title in {
            "四、核心功能范围",
            "七、创新点",
            "十、后续计划",
        }
        for text in paragraphs:
            paragraph = doc.add_paragraph(style="List Bullet" if use_bullets else None)
            paragraph.paragraph_format.space_after = Pt(2.5)
            paragraph.paragraph_format.line_spacing = 1.12
            paragraph.paragraph_format.first_line_indent = None if use_bullets else Cm(0.7)
            set_run_font(paragraph.add_run(text), 9.0)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, before, after in [
        ("Heading 1", 10.8, 5, 2),
        ("Heading 2", 9.8, 4, 2),
        ("Heading 3", 9.3, 3, 1),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT_HEX)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    set_run_font(title.add_run("MoonWheelKit 项目申报书"), 16.5, True, ACCENT_HEX)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(
        subtitle.add_run("分层时间轮与确定性虚拟时钟调度基础库"),
        9.3,
        False,
        "66737D",
    )

    h = doc.add_paragraph(style="Heading 1")
    set_run_font(h.add_run("一、基本信息"), 10.8, True, ACCENT_HEX)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    table_widths = [1700, 7900]
    for key, value in BASIC:
        cells = table.add_row().cells
        for cell, width in zip(cells, table_widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(cells[0], "EDF4F7")
        set_run_font(cells[0].paragraphs[0].add_run(key), 8.4, True, ACCENT_HEX)
        set_run_font(cells[1].paragraphs[0].add_run(value), 8.4)

    add_docx_sections(doc, PAGE1)
    doc.add_page_break()
    add_docx_sections(doc, PAGE2)

    for index, sec in enumerate(doc.sections):
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(footer.add_run("MoonWheelKit 项目申报书  |  共 2 页"), 7.5, False, "66737D")

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "MoonBit 国产开源生态大赛项目申报书"
    doc.core_properties.author = "MoonWheelKit contributors"
    doc.save(DOCX)


def build_text_files():
    lines = [f"# {TITLE}", "", "## 一、基本信息", ""]
    lines.extend([f"- {key}：{value}" for key, value in BASIC])
    for section_title, paragraphs in PAGE1 + PAGE2:
        lines.extend(["", f"## {section_title}", ""])
        lines.extend(paragraphs)
    MD.write_text("\n\n".join(lines), encoding="utf-8")
    CHECKLIST.write_text(
        "\n".join(
            [
                "# MoonWheelKit 报名检查清单",
                "",
                f"- [x] GitHub 仓库：{REPO}",
                "- [x] 15 次有效功能提交，超过 10 次要求",
                "- [x] 21 项确定性测试",
                "- [x] 默认、JS、Wasm、Wasm-GC 本地测试通过",
                "- [x] CLI、万级工作负载、CI、查重和路线图齐备",
                "- [x] 两页 PDF 与 DOCX 已生成",
                "- [ ] 填写参赛者电话、邮箱",
                "- [x] GitLink 仓库：https://www.gitlink.org.cn/brother666/MoonWheelKit",
                "- [x] GitHub Actions 首次运行成功：run 28453414903",
                "- [ ] GitLink 手动同步并核对有效提交数",
            ]
        ),
        encoding="utf-8",
    )


def copy_external():
    target = ROOT.parent / "MoonWheelKit_提交材料_最新版"
    target.mkdir(parents=True, exist_ok=True)
    for path in (PDF, DOCX, MD, CHECKLIST):
        shutil.copy2(path, target / path.name)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    build_text_files()
    copy_external()
    print(PDF)
    print(DOCX)
